import pandas as pd
import openai
from docx import Document


openai.api_key = ''

def gerar_docx(nome_medico, profissao, email, observacao_medico, nome_paciente, idade, tipo_sangue, doenca, observacao_paciente, nome_arquivo='output.docx'):
    doc = Document()
    doc.add_heading('Informações do Médico', level=1)
    doc.add_paragraph(f"Nome do medico: {nome_medico}")
    doc.add_paragraph(f"Profissao: {profissao}")
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Observacao: {observacao_medico}")
    doc.add_heading('Informações do Paciente', level=1)
    doc.add_paragraph(f"Nome paciente: {nome_paciente}")
    doc.add_paragraph(f"Idade: {idade}")
    doc.add_paragraph(f"Tipo de sangue: {tipo_sangue}")
    doc.add_paragraph(f"Doenca: {doenca}")
    doc.add_paragraph(f"Observacao: {observacao_paciente}")
    doc.save(nome_arquivo)

def gerar_excel(dados, nome_arquivo='output.xlsx'):
    df = pd.DataFrame(dados)
    df.to_excel(nome_arquivo, index=False)

def obter_resposta_chatgpt(prompt):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message['content'].strip()

def main():
   
    dados_medico = {
        'nome_medico': 'Dr. João Silva',
        'profissao': 'Cardiologista',
        'email': 'joao.silva@exemplo.com',
        'observacao_medico': 'Disponível para consultas às segundas e quartas.'
    }

    dados_paciente = {
        'nome_paciente': 'Maria Oliveira',
        'idade': 45,
        'tipo_sangue': 'O+',
        'doenca': 'Hipertensão',
        'observacao_paciente': 'Paciente deve realizar exames trimestrais.'
    }

    
    gerar_docx(
        dados_medico['nome_medico'],
        dados_medico['profissao'],
        dados_medico['email'],
        dados_medico['observacao_medico'],
        dados_paciente['nome_paciente'],
        dados_paciente['idade'],
        dados_paciente['tipo_sangue'],
        dados_paciente['doenca'],
        dados_paciente['observacao_paciente'],
        'relatorio_medico.docx'
    )

    dados_excel = [
        {
            'Nome do medico': dados_medico['nome_medico'],
            'Profissao': dados_medico['profissao'],
            'Email': dados_medico['email'],
            'Observacao': dados_medico['observacao_medico'],
            'Nome paciente': dados_paciente['nome_paciente'],
            'Idade': dados_paciente['idade'],
            'Tipo de sangue': dados_paciente['tipo_sangue'],
            'Doenca': dados_paciente['doenca'],
            'Observacao Paciente': dados_paciente['observacao_paciente']
        }
    ]

    gerar_excel(dados_excel, 'tabela_medica.xlsx')

    print("Arquivos gerados com sucesso!")

if __name__ == "__main__":
    main()
